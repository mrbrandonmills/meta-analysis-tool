"""APA 7th Edition Report Generator for Meta-Analysis.

This module provides comprehensive report generation capabilities with:
- APA 7th edition formatting
- Word (.docx) export
- PDF export
- Auto-generated citations and references
- Tables and figures
- Forest plots and funnel plots
"""

from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any
from io import BytesIO
import base64

from loguru import logger
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Image
from reportlab.lib import colors
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend


class APAFormatConfig:
    """Configuration for APA 7th edition formatting."""

    # Font settings
    FONT_NAME = "Times New Roman"
    FONT_SIZE = 12
    HEADING_SIZES = {
        1: 14,  # Centered, bold
        2: 14,  # Flush left, bold
        3: 14,  # Flush left, bold italic
        4: 12,  # Flush left, bold
        5: 12,  # Flush left, bold italic
    }

    # Spacing
    LINE_SPACING = 2.0  # Double spacing
    MARGIN_INCHES = 1.0

    # Page settings
    INCLUDE_RUNNING_HEAD = True
    INCLUDE_PAGE_NUMBERS = True

    # Abstract settings
    ABSTRACT_MAX_WORDS = 250


class APACitationFormatter:
    """Formats citations according to APA 7th edition."""

    @staticmethod
    def format_journal_article(
        authors: List[str],
        year: int,
        title: str,
        journal: str,
        volume: Optional[int] = None,
        issue: Optional[int] = None,
        pages: Optional[str] = None,
        doi: Optional[str] = None,
    ) -> str:
        """Format a journal article citation.

        Example:
            Smith, J. D., & Jones, A. B. (2020). Effects of mindfulness on anxiety:
            A meta-analysis. Journal of Clinical Psychology, 76(5), 123-145.
            https://doi.org/10.1234/jcp.2020.12345
        """
        # Format authors
        formatted_authors = APACitationFormatter._format_author_list(authors)

        # Build citation
        citation = f"{formatted_authors} ({year}). {title}. {journal}"

        if volume:
            citation += f", {volume}"
        if issue:
            citation += f"({issue})"
        if pages:
            citation += f", {pages}"
        citation += "."

        if doi:
            citation += f" https://doi.org/{doi}"

        return citation

    @staticmethod
    def _format_author_list(authors: List[str]) -> str:
        """Format author list according to APA style.

        Rules:
        - 1-2 authors: List all (Smith, J. D., & Jones, A. B.)
        - 3-20 authors: List all
        - 21+ authors: List first 19, then ... then final author
        """
        if not authors:
            return ""

        if len(authors) == 1:
            return authors[0]

        if len(authors) == 2:
            return f"{authors[0]}, & {authors[1]}"

        if len(authors) <= 20:
            return ", ".join(authors[:-1]) + f", & {authors[-1]}"

        # 21+ authors
        first_19 = ", ".join(authors[:19])
        return f"{first_19}, ... {authors[-1]}"

    @staticmethod
    def format_in_text_citation(authors: List[str], year: int) -> str:
        """Format an in-text citation.

        Examples:
            (Smith, 2020)
            (Smith & Jones, 2020)
            (Smith et al., 2020)
        """
        if not authors:
            return f"(Unknown, {year})"

        if len(authors) == 1:
            return f"({authors[0]}, {year})"

        if len(authors) == 2:
            return f"({authors[0]} & {authors[1]}, {year})"

        # 3+ authors
        return f"({authors[0]} et al., {year})"


class APAReportGenerator:
    """Main class for generating APA-formatted meta-analysis reports."""

    def __init__(self, output_dir: Optional[Path] = None):
        """Initialize report generator.

        Args:
            output_dir: Directory for saving generated reports
        """
        self.output_dir = output_dir or Path("/tmp/reports")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.citation_formatter = APACitationFormatter()
        logger.info(f"Initialized APAReportGenerator with output_dir: {self.output_dir}")

    def generate_report(
        self,
        analysis_data: Dict[str, Any],
        format: str = "docx",
        custom_sections: Optional[Dict[str, str]] = None,
    ) -> Dict[str, Any]:
        """Generate a complete APA-formatted report.

        Args:
            analysis_data: Meta-analysis data including studies, results, etc.
            format: Output format ('docx', 'pdf', or 'both')
            custom_sections: Optional custom content for sections

        Returns:
            Dict with file paths and metadata
        """
        logger.info(f"Generating {format} report for analysis {analysis_data.get('id', 'unknown')}")

        results = {}

        if format in ("docx", "both"):
            docx_path = self._generate_word_document(analysis_data, custom_sections)
            results["docx_path"] = str(docx_path)

        if format in ("pdf", "both"):
            pdf_path = self._generate_pdf_document(analysis_data, custom_sections)
            results["pdf_path"] = str(pdf_path)

        results["generated_at"] = datetime.utcnow().isoformat()
        results["format"] = format

        logger.info(f"Report generation complete: {results}")
        return results

    def _generate_word_document(
        self,
        analysis_data: Dict[str, Any],
        custom_sections: Optional[Dict[str, str]] = None,
    ) -> Path:
        """Generate Word document with APA formatting."""
        logger.info("Generating Word document")

        doc = Document()
        self._apply_apa_styles(doc)

        # Title Page
        self._add_title_page(doc, analysis_data)
        doc.add_page_break()

        # Abstract
        self._add_abstract(doc, analysis_data, custom_sections)
        doc.add_page_break()

        # Introduction
        self._add_introduction(doc, analysis_data, custom_sections)

        # Methods
        self._add_methods(doc, analysis_data, custom_sections)
        doc.add_page_break()

        # Results
        self._add_results(doc, analysis_data, custom_sections)
        doc.add_page_break()

        # Discussion
        self._add_discussion(doc, analysis_data, custom_sections)
        doc.add_page_break()

        # References
        self._add_references(doc, analysis_data)

        # Save document
        analysis_id = analysis_data.get("id", "unknown")
        filename = f"meta_analysis_report_{analysis_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = self.output_dir / filename
        doc.save(output_path)

        logger.info(f"Word document saved to {output_path}")
        return output_path

    def _apply_apa_styles(self, doc: Document) -> None:
        """Apply APA 7th edition styles to document."""
        # Set document defaults
        style = doc.styles['Normal']
        font = style.font
        font.name = APAFormatConfig.FONT_NAME
        font.size = Pt(APAFormatConfig.FONT_SIZE)

        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = APAFormatConfig.LINE_SPACING
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

        # Create heading styles
        for level in range(1, 6):
            style_name = f'Heading {level}'
            if style_name in doc.styles:
                heading_style = doc.styles[style_name]
            else:
                heading_style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

            heading_font = heading_style.font
            heading_font.name = APAFormatConfig.FONT_NAME
            heading_font.size = Pt(APAFormatConfig.HEADING_SIZES[level])
            heading_font.bold = True

            if level in (3, 5):
                heading_font.italic = True

            heading_paragraph = heading_style.paragraph_format
            heading_paragraph.line_spacing = APAFormatConfig.LINE_SPACING

            if level == 1:
                heading_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_title_page(self, doc: Document, analysis_data: Dict[str, Any]) -> None:
        """Add APA-formatted title page."""
        # Running head
        if APAFormatConfig.INCLUDE_RUNNING_HEAD:
            title = analysis_data.get("title", "META-ANALYSIS")
            running_head = title[:50].upper()
            p = doc.add_paragraph(f"Running head: {running_head}")
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Add vertical space (approximately 1/3 down the page)
        for _ in range(8):
            doc.add_paragraph()

        # Title (bold, centered)
        title = analysis_data.get("title", "Meta-Analysis of [Research Question]")
        p = doc.add_paragraph(title)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph()

        # Author(s)
        authors = analysis_data.get("authors", ["Author Name"])
        for author in authors:
            p = doc.add_paragraph(author)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Institution
        institution = analysis_data.get("institution", "Institution Name")
        p = doc.add_paragraph(institution)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add vertical space
        for _ in range(8):
            doc.add_paragraph()

        # Author Note (optional)
        if "author_note" in analysis_data:
            p = doc.add_paragraph("Author Note")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True

            doc.add_paragraph()
            p = doc.add_paragraph(analysis_data["author_note"])
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_abstract(
        self,
        doc: Document,
        analysis_data: Dict[str, Any],
        custom_sections: Optional[Dict[str, str]] = None,
    ) -> None:
        """Add abstract section."""
        # Abstract heading (centered, bold)
        p = doc.add_paragraph("Abstract")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True

        doc.add_paragraph()

        # Abstract text
        if custom_sections and "abstract" in custom_sections:
            abstract_text = custom_sections["abstract"]
        else:
            abstract_text = self._generate_abstract(analysis_data)

        doc.add_paragraph(abstract_text)

        doc.add_paragraph()

        # Keywords
        keywords = analysis_data.get("keywords", ["meta-analysis", "systematic review"])
        p = doc.add_paragraph()
        run = p.add_run("Keywords: ")
        run.italic = True
        p.add_run(", ".join(keywords))

    def _generate_abstract(self, analysis_data: Dict[str, Any]) -> str:
        """Auto-generate abstract from analysis data."""
        research_question = analysis_data.get("research_question", "the research question")
        num_studies = analysis_data.get("num_studies", 0)
        num_participants = analysis_data.get("num_participants", 0)

        effect_size = analysis_data.get("pooled_effect_size", 0.0)
        ci_lower = analysis_data.get("ci_lower", 0.0)
        ci_upper = analysis_data.get("ci_upper", 0.0)

        abstract = (
            f"This meta-analysis examined {research_question}. "
            f"A systematic search identified {num_studies} studies "
            f"with a total of {num_participants} participants. "
            f"The pooled effect size was {effect_size:.3f} "
            f"(95% CI [{ci_lower:.3f}, {ci_upper:.3f}]). "
            f"These findings suggest significant effects. "
            f"Implications for theory and practice are discussed."
        )

        return abstract

    def _add_introduction(
        self,
        doc: Document,
        analysis_data: Dict[str, Any],
        custom_sections: Optional[Dict[str, str]] = None,
    ) -> None:
        """Add introduction section."""
        # Section heading (centered, bold)
        p = doc.add_paragraph("Introduction")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True

        doc.add_paragraph()

        # Introduction text
        if custom_sections and "introduction" in custom_sections:
            intro_text = custom_sections["introduction"]
        else:
            intro_text = self._generate_introduction(analysis_data)

        doc.add_paragraph(intro_text)

        # Research question subsection
        doc.add_paragraph()
        p = doc.add_paragraph("Research Question")
        run = p.runs[0]
        run.bold = True

        research_question = analysis_data.get("research_question", "What is the effect of X on Y?")
        doc.add_paragraph(f"    {research_question}")

    def _generate_introduction(self, analysis_data: Dict[str, Any]) -> str:
        """Auto-generate introduction text."""
        topic = analysis_data.get("topic", "the research topic")

        intro = (
            f"    {topic.capitalize()} has been the subject of considerable research "
            f"in recent years. However, findings across individual studies have been "
            f"inconsistent, with some reporting significant effects while others found "
            f"null or contradictory results. This meta-analysis aims to synthesize "
            f"the existing literature to provide a comprehensive assessment of the "
            f"evidence. By pooling data from multiple studies, we can obtain more "
            f"precise estimates of effect sizes and identify potential moderating factors."
        )

        return intro

    def _add_methods(
        self,
        doc: Document,
        analysis_data: Dict[str, Any],
        custom_sections: Optional[Dict[str, str]] = None,
    ) -> None:
        """Add methods section."""
        p = doc.add_paragraph("Methods")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True

        doc.add_paragraph()

        # Search Strategy
        p = doc.add_paragraph("Search Strategy")
        run = p.runs[0]
        run.bold = True

        databases = analysis_data.get("databases", ["PubMed", "PsycINFO"])
        search_terms = analysis_data.get("search_terms", [])

        search_text = (
            f"    A systematic literature search was conducted using the following databases: "
            f"{', '.join(databases)}. Search terms included: {', '.join(search_terms)}. "
            f"The search was conducted on {analysis_data.get('search_date', datetime.now().strftime('%B %d, %Y'))}."
        )
        doc.add_paragraph(search_text)

        doc.add_paragraph()

        # Inclusion/Exclusion Criteria
        p = doc.add_paragraph("Inclusion and Exclusion Criteria")
        run = p.runs[0]
        run.bold = True

        doc.add_paragraph("    Studies were included if they met the following criteria:")

        inclusion = analysis_data.get("inclusion_criteria", [
            "Published in peer-reviewed journals",
            "Randomized controlled trials",
            "Adult participants"
        ])
        for criterion in inclusion:
            doc.add_paragraph(f"    • {criterion}")

        doc.add_paragraph()
        doc.add_paragraph("    Studies were excluded if they:")

        exclusion = analysis_data.get("exclusion_criteria", [
            "Were not available in English",
            "Did not report sufficient statistical data"
        ])
        for criterion in exclusion:
            doc.add_paragraph(f"    • {criterion}")

        doc.add_paragraph()

        # Statistical Analysis
        p = doc.add_paragraph("Statistical Analysis")
        run = p.runs[0]
        run.bold = True

        analysis_method = analysis_data.get("analysis_method", "random-effects model")
        doc.add_paragraph(
            f"    Meta-analyses were conducted using a {analysis_method}. "
            f"Effect sizes were calculated as standardized mean differences. "
            f"Heterogeneity was assessed using the I² statistic and Q test. "
            f"Publication bias was evaluated using funnel plots and Egger's test."
        )

    def _add_results(
        self,
        doc: Document,
        analysis_data: Dict[str, Any],
        custom_sections: Optional[Dict[str, str]] = None,
    ) -> None:
        """Add results section with tables and figures."""
        p = doc.add_paragraph("Results")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True

        doc.add_paragraph()

        # Study Selection
        p = doc.add_paragraph("Study Selection")
        run = p.runs[0]
        run.bold = True

        num_identified = analysis_data.get("num_identified", 0)
        num_screened = analysis_data.get("num_screened", 0)
        num_included = analysis_data.get("num_studies", 0)

        doc.add_paragraph(
            f"    The database search identified {num_identified} records. "
            f"After removing duplicates and screening titles and abstracts, "
            f"{num_screened} articles were assessed for eligibility. "
            f"A total of {num_included} studies met the inclusion criteria "
            f"and were included in the meta-analysis."
        )

        doc.add_paragraph()

        # Study Characteristics
        p = doc.add_paragraph("Study Characteristics")
        run = p.runs[0]
        run.bold = True

        doc.add_paragraph(
            f"    Table 1 presents the characteristics of included studies. "
            f"Studies were published between {analysis_data.get('year_range', '2010-2024')}."
        )

        # Add study characteristics table
        self._add_study_table(doc, analysis_data)

        doc.add_paragraph()

        # Meta-Analysis Results
        p = doc.add_paragraph("Meta-Analysis Results")
        run = p.runs[0]
        run.bold = True

        effect_size = analysis_data.get("pooled_effect_size", 0.0)
        ci_lower = analysis_data.get("ci_lower", 0.0)
        ci_upper = analysis_data.get("ci_upper", 0.0)
        p_value = analysis_data.get("p_value", 0.05)
        i_squared = analysis_data.get("i_squared", 0.0)

        doc.add_paragraph(
            f"    The pooled effect size was {effect_size:.3f} "
            f"(95% CI [{ci_lower:.3f}, {ci_upper:.3f}], p = {p_value:.3f}). "
            f"Heterogeneity was {'substantial' if i_squared > 50 else 'low'} "
            f"(I² = {i_squared:.1f}%)."
        )

        doc.add_paragraph()
        doc.add_paragraph("    See Figure 1 for the forest plot of individual study effects.")

    def _add_study_table(self, doc: Document, analysis_data: Dict[str, Any]) -> None:
        """Add table of study characteristics."""
        doc.add_paragraph()

        # Table caption
        p = doc.add_paragraph("Table 1")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.italic = True

        p = doc.add_paragraph("Characteristics of Included Studies")
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.italic = True

        # Create table
        studies = analysis_data.get("studies", [])
        if studies:
            table = doc.add_table(rows=len(studies) + 1, cols=5)
            table.style = 'Light Grid Accent 1'

            # Header row
            headers = ['Study', 'N', 'Design', 'Effect Size', 'Quality']
            for i, header in enumerate(headers):
                cell = table.rows[0].cells[i]
                cell.text = header
                run = cell.paragraphs[0].runs[0]
                run.bold = True

            # Data rows
            for i, study in enumerate(studies[:10], 1):  # Limit to 10 studies for example
                table.rows[i].cells[0].text = study.get("authors", ["Unknown"])[0] + " et al."
                table.rows[i].cells[1].text = str(study.get("sample_size", "-"))
                table.rows[i].cells[2].text = study.get("design", "RCT")
                table.rows[i].cells[3].text = f"{study.get('effect_size', 0.0):.2f}"
                table.rows[i].cells[4].text = study.get("quality_rating", "Moderate")

    def _add_discussion(
        self,
        doc: Document,
        analysis_data: Dict[str, Any],
        custom_sections: Optional[Dict[str, str]] = None,
    ) -> None:
        """Add discussion section."""
        p = doc.add_paragraph("Discussion")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True

        doc.add_paragraph()

        # Main findings
        if custom_sections and "discussion" in custom_sections:
            discussion_text = custom_sections["discussion"]
        else:
            discussion_text = self._generate_discussion(analysis_data)

        doc.add_paragraph(discussion_text)

        doc.add_paragraph()

        # Limitations
        p = doc.add_paragraph("Limitations")
        run = p.runs[0]
        run.bold = True

        limitations = analysis_data.get("limitations", [
            "Limited to English-language publications",
            "Heterogeneity across studies suggests variability in effects",
            "Potential publication bias cannot be ruled out"
        ])

        doc.add_paragraph("    This meta-analysis has several limitations:")
        for limitation in limitations:
            doc.add_paragraph(f"    • {limitation}")

        doc.add_paragraph()

        # Conclusions
        p = doc.add_paragraph("Conclusions")
        run = p.runs[0]
        run.bold = True

        doc.add_paragraph(
            f"    This meta-analysis provides strong evidence for "
            f"{analysis_data.get('research_question', 'the research question')}. "
            f"The findings have important implications for practice and future research."
        )

    def _generate_discussion(self, analysis_data: Dict[str, Any]) -> str:
        """Auto-generate discussion text."""
        effect_size = analysis_data.get("pooled_effect_size", 0.0)

        magnitude = "large" if abs(effect_size) > 0.8 else "moderate" if abs(effect_size) > 0.5 else "small"

        discussion = (
            f"    The present meta-analysis found a {magnitude} effect size "
            f"({effect_size:.3f}), indicating that the intervention has a meaningful "
            f"impact. This finding is consistent with theoretical predictions and "
            f"extends previous research by providing a quantitative synthesis of "
            f"the evidence. The results suggest that practitioners can confidently "
            f"implement these interventions with expectation of beneficial outcomes."
        )

        return discussion

    def _add_references(self, doc: Document, analysis_data: Dict[str, Any]) -> None:
        """Add references section with APA-formatted citations."""
        p = doc.add_paragraph("References")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True

        doc.add_paragraph()

        studies = analysis_data.get("studies", [])

        # Sort studies alphabetically by first author
        sorted_studies = sorted(
            studies,
            key=lambda s: s.get("authors", ["Unknown"])[0].split(",")[0]
        )

        for study in sorted_studies:
            citation = self.citation_formatter.format_journal_article(
                authors=study.get("authors", ["Unknown"]),
                year=study.get("year", 2020),
                title=study.get("title", "Unknown title"),
                journal=study.get("journal", "Unknown journal"),
                volume=study.get("volume"),
                issue=study.get("issue"),
                pages=study.get("pages"),
                doi=study.get("doi"),
            )

            # APA hanging indent
            p = doc.add_paragraph(citation)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)

    def _generate_pdf_document(
        self,
        analysis_data: Dict[str, Any],
        custom_sections: Optional[Dict[str, str]] = None,
    ) -> Path:
        """Generate PDF document with APA formatting."""
        logger.info("Generating PDF document")

        analysis_id = analysis_data.get("id", "unknown")
        filename = f"meta_analysis_report_{analysis_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
        output_path = self.output_dir / filename

        # Create PDF document
        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=letter,
            leftMargin=APAFormatConfig.MARGIN_INCHES * inch,
            rightMargin=APAFormatConfig.MARGIN_INCHES * inch,
            topMargin=APAFormatConfig.MARGIN_INCHES * inch,
            bottomMargin=APAFormatConfig.MARGIN_INCHES * inch,
        )

        # Build styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='APA',
            fontName='Times-Roman',
            fontSize=12,
            leading=24,  # Double spacing (2 * fontSize)
            alignment=0,  # Left aligned
        ))
        styles.add(ParagraphStyle(
            name='APATitle',
            fontName='Times-Bold',
            fontSize=14,
            leading=24,
            alignment=1,  # Centered
        ))
        styles.add(ParagraphStyle(
            name='APAHeading',
            fontName='Times-Bold',
            fontSize=14,
            leading=24,
            alignment=1,  # Centered
        ))

        # Build content
        story = []

        # Title page
        title = analysis_data.get("title", "Meta-Analysis Report")
        story.append(Paragraph(title, styles['APATitle']))
        story.append(Spacer(1, 0.5 * inch))

        authors = analysis_data.get("authors", ["Author Name"])
        for author in authors:
            story.append(Paragraph(author, styles['APA']))

        story.append(Spacer(1, 0.3 * inch))
        institution = analysis_data.get("institution", "Institution Name")
        story.append(Paragraph(institution, styles['APA']))
        story.append(PageBreak())

        # Abstract
        story.append(Paragraph("Abstract", styles['APAHeading']))
        story.append(Spacer(1, 0.2 * inch))

        abstract_text = (
            custom_sections.get("abstract") if custom_sections
            else self._generate_abstract(analysis_data)
        )
        story.append(Paragraph(abstract_text, styles['APA']))
        story.append(PageBreak())

        # Introduction
        story.append(Paragraph("Introduction", styles['APAHeading']))
        story.append(Spacer(1, 0.2 * inch))

        intro_text = (
            custom_sections.get("introduction") if custom_sections
            else self._generate_introduction(analysis_data)
        )
        story.append(Paragraph(intro_text, styles['APA']))
        story.append(PageBreak())

        # Build PDF
        doc.build(story)

        logger.info(f"PDF document saved to {output_path}")
        return output_path

    def generate_forest_plot(
        self,
        studies: List[Dict[str, Any]],
        output_path: Optional[Path] = None,
    ) -> Path:
        """Generate forest plot visualization.

        Args:
            studies: List of study data with effect sizes and confidence intervals
            output_path: Optional path to save the plot

        Returns:
            Path to saved plot image
        """
        if output_path is None:
            output_path = self.output_dir / f"forest_plot_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"

        fig, ax = plt.subplots(figsize=(10, max(6, len(studies) * 0.5)))

        # Extract data
        study_names = [s.get("authors", ["Unknown"])[0] + " et al." for s in studies]
        effect_sizes = [s.get("effect_size", 0.0) for s in studies]
        ci_lower = [s.get("ci_lower", 0.0) for s in studies]
        ci_upper = [s.get("ci_upper", 0.0) for s in studies]

        # Plot individual studies
        y_positions = range(len(studies))
        ax.errorbar(
            effect_sizes,
            y_positions,
            xerr=[[es - ci_l for es, ci_l in zip(effect_sizes, ci_lower)],
                  [ci_u - es for es, ci_u in zip(effect_sizes, ci_upper)]],
            fmt='o',
            color='black',
            ecolor='gray',
            capsize=5,
        )

        # Add vertical line at 0
        ax.axvline(x=0, color='black', linestyle='--', linewidth=0.8)

        # Labels
        ax.set_yticks(y_positions)
        ax.set_yticklabels(study_names)
        ax.set_xlabel('Effect Size (Standardized Mean Difference)')
        ax.set_title('Forest Plot of Study Effect Sizes')
        ax.grid(axis='x', alpha=0.3)

        plt.tight_layout()
        plt.savefig(output_path, dpi=300, bbox_inches='tight')
        plt.close()

        logger.info(f"Forest plot saved to {output_path}")
        return output_path

    def generate_funnel_plot(
        self,
        studies: List[Dict[str, Any]],
        output_path: Optional[Path] = None,
    ) -> Path:
        """Generate funnel plot for publication bias assessment.

        Args:
            studies: List of study data with effect sizes and standard errors
            output_path: Optional path to save the plot

        Returns:
            Path to saved plot image
        """
        if output_path is None:
            output_path = self.output_dir / f"funnel_plot_{datetime.now().strftime('%Y%m%d_%H%M%S')}.png"

        fig, ax = plt.subplots(figsize=(8, 6))

        # Extract data
        effect_sizes = [s.get("effect_size", 0.0) for s in studies]
        standard_errors = [s.get("standard_error", 0.1) for s in studies]

        # Plot studies
        ax.scatter(effect_sizes, standard_errors, alpha=0.6, edgecolors='black')

        # Add reference line (pooled effect)
        pooled_effect = sum(effect_sizes) / len(effect_sizes) if effect_sizes else 0
        ax.axvline(x=pooled_effect, color='black', linestyle='--', linewidth=1)

        # Invert y-axis (larger studies at top)
        ax.invert_yaxis()

        # Labels
        ax.set_xlabel('Effect Size')
        ax.set_ylabel('Standard Error')
        ax.set_title('Funnel Plot for Publication Bias Assessment')
        ax.grid(alpha=0.3)

        plt.tight_layout()
        plt.savefig(output_path, dpi=300, bbox_inches='tight')
        plt.close()

        logger.info(f"Funnel plot saved to {output_path}")
        return output_path
